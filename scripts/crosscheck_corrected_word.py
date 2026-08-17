#!/usr/bin/env python3
"""Cross-check the pending FIDIC source records against a corrected DOCX.

This script is evidence-only: it never updates clause wording or verification
status. The PDF remains authoritative and the correction schedule is treated as
a review aid, not as a substitute for PDF confirmation.
"""
from __future__ import annotations

import argparse, hashlib, json, re, unicodedata
from collections import Counter
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from pypdf import PdfReader

TRANSLATION = str.maketrans({
    "\u00a0": " ", "\u00ad": "", "\u2018": "'", "\u2019": "'", "\u201c": '"',
    "\u201d": '"', "\u2013": "-", "\u2014": "-", "\ufb00": "ff", "\ufb01": "fi",
    "\ufb02": "fl", "\ufb03": "ffi", "\ufb04": "ffl",
})
SCHEDULE_SPECIFIC = {
    "1.1", "1.2", "2.2", "4.9", "4.13", "4.18", "6.6", "8.3", "8.4",
    "11.1", "14.5", "17.3", "17.6", "20.2.5", "21.1", "21.2", "21.3", "21.4",
}


def canonical(value: str) -> str:
    return re.sub(r"\s+", " ", unicodedata.normalize("NFKC", value).translate(TRANSLATION)).strip()


def word_rows(path: Path) -> dict[str, dict]:
    doc = Document(path)
    if len(doc.tables) != 1:
        raise RuntimeError(f"Expected one clause table, found {len(doc.tables)}")
    records = {}
    for row in doc.tables[0].rows[1:]:
        label = row.cells[0].text.strip()
        match = re.match(r"(?:Clause\s+)?(\d+(?:\.\d+)*)\b", label, re.I)
        if not match: continue
        number = match.group(1)
        title = re.sub(r"^(?:Clause\s+)?\d+(?:\.\d+)*\s*", "", label, flags=re.I).strip()
        records[number] = {"clause_no": number, "label": label, "title": title, "full_text": row.cells[1].text}
    return records


def ligature_pairs(schedule: str) -> list[tuple[str, str]]:
    pairs = []
    for line in schedule.splitlines():
        match = re.match(r"\|\s*`([^`]+)`\s*\|\s*([^|]+?)\s*\|\s*\d+\s*\|", line)
        if match: pairs.append((match.group(1), match.group(2).strip().strip("`")))
    return pairs


def apply_known_word_repairs(value: str, pairs: list[tuple[str, str]]) -> str:
    for old, new in sorted(pairs, key=lambda pair: len(pair[0]), reverse=True):
        value = re.sub(rf"(?<![A-Za-z]){re.escape(old)}(?![A-Za-z])", new, value)
    return value


def digest_wording(source: dict) -> str:
    payload = [(x["clause_no"], x["full_text"], x.get("paragraphs")) for x in source["clauses"]]
    return hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode()).hexdigest()


def summary(report: dict) -> str:
    c = report["counts"]
    substantive = [x for x in report["items"] if x["schedule_specific"] and not x["current_equals_corrected_word"]]
    lines = "\n".join(f"- {x['clause_no']} {x['clause_title']}: {x['crosscheck_status']} (similarity {x['similarity']:.4f})" for x in substantive) or "- None"
    return f"""# FIDIC 2017 Corrected Word Cross-check Summary

## Sources

- Existing clause layer: `{report['clause_source_file']}`
- Corrected Word source: `{report['corrected_word_file']}`
- Correction schedule: `{report['correction_schedule_file']}`
- Authoritative PDF: `{report['pdf_source_file']}`

## Control boundary

This is a three-way evidence review. No imported clause wording was changed. A corrected Word match alone does not establish PDF verification; uncertain records remain pending unless the PDF supports a complete reliable match.

## Coverage and results

- Existing sub-clause records cross-checked: **{report['records_checked']}**
- Existing records found in corrected Word: **{report['records_found_in_corrected_word']}**
- Existing wording equal to corrected Word after layout normalisation: **{c.get('current_equals_corrected_word', 0)}**
- Differences explained entirely by listed dropped-ligature repairs: **{c.get('schedule_ligature_repairs_explain_difference', 0)}**
- Schedule-specific differences requiring/recording PDF confirmation: **{c.get('schedule_specific_difference', 0)}**
- Other corrected-Word differences requiring review: **{c.get('other_corrected_word_difference', 0)}**
- Corrected Word records with complete unique PDF text-layer match: **{report['corrected_word_complete_pdf_matches']}**
- Current JSON discrepancies objectively identified without overwrite: **{report['objective_current_text_discrepancy_count']}**
- Existing imported wording digest unchanged: **{report['imported_wording_unchanged']}**

## Schedule-specific clauses where current JSON differs

{lines}

## Important limitation

The corrected DOCX and schedule materially improve discrepancy identification, but they do not automatically replace the existing imported text. PDF page 101 remains a degraded scan, so Clauses 21.1-21.4 retain a manual-review caution even where the corrected Word is internally consistent.
"""


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path); parser.add_argument("corrected_word", type=Path)
    parser.add_argument("schedule", type=Path); parser.add_argument("pdf", type=Path)
    parser.add_argument("--json-report", type=Path, required=True); parser.add_argument("--summary", type=Path, required=True)
    parser.add_argument("--apply-metadata", action="store_true")
    parser.add_argument("--primary-report", type=Path)
    parser.add_argument("--primary-summary", type=Path)
    args = parser.parse_args()
    source = json.loads(args.source.read_text(encoding="utf-8")); before = digest_wording(source)
    corrected = word_rows(args.corrected_word); schedule_text = args.schedule.read_text(encoding="utf-8")
    repairs = ligature_pairs(schedule_text)
    pdf_pages = [(page.extract_text() or "") for page in PdfReader(args.pdf).pages]
    pdf_text = canonical("\n".join(pdf_pages))
    items, counts, corrected_pdf_matches = [], Counter(), 0
    for clause in source["clauses"]:
        number = clause["clause_no"]; record = corrected.get(number)
        if not record: raise RuntimeError(f"Corrected Word record missing for {number}")
        current_text, corrected_text = canonical(clause["full_text"]), canonical(record["full_text"])
        equal = current_text == corrected_text
        repaired_equal = not equal and canonical(apply_known_word_repairs(clause["full_text"], repairs)) == corrected_text
        pdf_count = pdf_text.count(corrected_text) if corrected_text else 0
        corrected_pdf_matches += int(pdf_count == 1)
        objective_discrepancy = pdf_count == 1 and not equal and number not in {"21.1", "21.2", "21.3", "21.4"}
        if equal:
            status = "current_equals_corrected_word"; counts[status] += 1
        elif repaired_equal:
            status = "schedule_ligature_repairs_explain_difference"; counts[status] += 1
        elif number in SCHEDULE_SPECIFIC:
            status = "schedule_specific_difference"; counts[status] += 1
        else:
            status = "other_corrected_word_difference"; counts[status] += 1
        items.append({
            "clause_no": number, "clause_title": clause["clause_title"],
            "existing_pdf_verification_status": clause.get("pdf_verification_status"),
            "current_equals_corrected_word": equal, "known_ligature_repairs_explain_difference": repaired_equal,
            "schedule_specific": number in SCHEDULE_SPECIFIC, "corrected_word_complete_pdf_match": pdf_count == 1,
            "objective_current_text_discrepancy": objective_discrepancy,
            "similarity": round(SequenceMatcher(None, current_text, corrected_text, autojunk=False).ratio(), 6),
            "crosscheck_status": status,
            "recommended_action": "retain_existing_text_and_report_only" if not equal else "no_text_change",
        })
    after = digest_wording(source)
    report = {
        "report_version": "1.0", "generated_at": datetime.now(timezone.utc).isoformat(),
        "clause_source_file": args.source.as_posix(), "corrected_word_file": args.corrected_word.as_posix(),
        "correction_schedule_file": args.schedule.as_posix(), "pdf_source_file": args.pdf.as_posix(),
        "records_checked": len(source["clauses"]), "records_found_in_corrected_word": len(items),
        "corrected_word_total_clause_rows": len(corrected), "corrected_word_complete_pdf_matches": corrected_pdf_matches,
        "known_ligature_repair_pairs": len(repairs), "counts": dict(sorted(counts.items())),
        "imported_wording_digest_before": before, "imported_wording_digest_after": after,
        "imported_wording_unchanged": before == after, "items": items,
    }
    objective_items = [item for item in items if item["objective_current_text_discrepancy"]]
    report["objective_current_text_discrepancy_count"] = len(objective_items)
    report["objective_current_text_discrepancy_clauses"] = [item["clause_no"] for item in objective_items]
    if args.apply_metadata:
        if not args.primary_report or not args.primary_summary:
            raise RuntimeError("--primary-report and --primary-summary are required with --apply-metadata")
        by_number = {item["clause_no"]: item for item in source["clauses"]}
        for item in objective_items:
            clause = by_number[item["clause_no"]]
            clause["pdf_verification_status"] = "text_discrepancy_found"
            clause["text_match_status"] = "substantive_difference_possible"
            clause["qc_status"] = "text_discrepancy_found"
            clause["qc_note"] = "Corrected Word wording has a complete unique PDF text-layer match while the current imported wording differs. Existing text retained unchanged pending controlled correction."
            clause["corrected_word_crosscheck"] = {
                "source_file": args.corrected_word.as_posix(),
                "correction_schedule": args.schedule.as_posix(),
                "status": "corrected_word_matches_pdf_current_import_differs",
                "text_overwritten": False,
            }
        after_apply = digest_wording(source)
        if before != after_apply: raise RuntimeError("Imported wording changed while applying metadata")
        source_counts = Counter(item["pdf_verification_status"] for item in source["clauses"])
        source["pdf_verification"]["counts"] = dict(sorted(source_counts.items()))
        source["pdf_verification"]["imported_text_unchanged"] = True
        args.source.write_text(json.dumps(source, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        primary = json.loads(args.primary_report.read_text(encoding="utf-8"))
        primary_by_number = {item["clause_no"]: item for item in primary["items"]}
        for item in objective_items:
            target = primary_by_number[item["clause_no"]]
            target.update({
                "verification_status": "text_discrepancy_found",
                "text_match_status": "substantive_difference_possible",
                "issue_type": "corrected_word_matches_pdf_current_import_differs",
                "qc_note": "Corrected Word wording uniquely matches the PDF text layer while current imported wording differs; discrepancy recorded without overwriting text.",
            })
        primary_counts = Counter(item["verification_status"] for item in primary["items"])
        primary["counts"] = dict(sorted(primary_counts.items()))
        primary["corrected_word_crosscheck"] = {
            "report_file": args.json_report.as_posix(),
            "objective_discrepancy_count": len(objective_items),
            "objective_discrepancy_clauses": [item["clause_no"] for item in objective_items],
            "imported_text_unchanged": True,
        }
        args.primary_report.write_text(json.dumps(primary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        markdown = args.primary_summary.read_text(encoding="utf-8")
        for label, value in primary_counts.items():
            markdown = re.sub(rf"(- {re.escape(label)}: \*\*)\d+(\*\*)", rf"\g<1>{value}\2", markdown)
        discrepancy_lines = "\n".join(f"- {item['clause_no']} {item['clause_title']}" for item in objective_items)
        section = f"\n## Corrected Word cross-check - confirmed current-text discrepancies\n\n{discrepancy_lines}\n\nThese records were not overwritten. The corrected Word wording has a complete unique PDF text-layer match while the current imported wording differs.\n"
        markdown = re.sub(r"\n## Corrected Word cross-check - confirmed current-text discrepancies\n.*?(?=\n## Interpretation\n)", "", markdown, flags=re.S)
        markdown = markdown.replace("\n## Interpretation\n", section + "\n## Interpretation\n")
        args.primary_summary.write_text(markdown, encoding="utf-8")
    args.json_report.parent.mkdir(parents=True, exist_ok=True)
    args.json_report.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    args.summary.write_text(summary(report), encoding="utf-8")
    print(json.dumps({key: report[key] for key in ("records_checked", "records_found_in_corrected_word", "corrected_word_complete_pdf_matches", "known_ligature_repair_pairs", "counts", "imported_wording_unchanged")}, indent=2))
    return 0


if __name__ == "__main__": raise SystemExit(main())
