#!/usr/bin/env python3
"""Regenerate the FIDIC 2017 Clause Spine source layer from corrected DOCX."""
from __future__ import annotations

import argparse, hashlib, json, re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from import_fidic_word import (
    literal_cross_references,
    make_clause_id,
    paragraph_records,
    validate_source_layer,
)

SOURCE_ORIGIN = "Corrected Word working copy"
ROW_PATTERN = re.compile(r"^(?:Clause\s+)?(\d+(?:\.\d+)*)\b(.*)$", re.I | re.S)


def clean_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    lines = [line.rstrip() for line in value.split("\n")]
    while lines and not lines[0].strip(): lines.pop(0)
    while lines and not lines[-1].strip(): lines.pop()
    return "\n".join(lines)


def parse_docx(path: Path) -> tuple[list[dict], list[dict], dict]:
    document = Document(path)
    if len(document.tables) != 1:
        raise RuntimeError(f"Expected exactly one clause table, found {len(document.tables)}")
    table = document.tables[0]
    if len(table.columns) != 2:
        raise RuntimeError(f"Expected a two-column clause table, found {len(table.columns)} columns")
    main_clauses, clauses, issues, omitted_empty_containers = [], [], [], []
    current_parent_no = current_parent_title = None
    for row_index, row in enumerate(table.rows[1:], start=1):
        left, right = clean_text(row.cells[0].text).strip(), clean_text(row.cells[1].text)
        match = ROW_PATTERN.match(left)
        if not match:
            issues.append(f"Row {row_index}: unrecognised clause label {left[:80]!r}")
            continue
        number, label_tail = match.group(1), clean_text(match.group(2))
        if left.lower().startswith("clause ") and "." not in number:
            current_parent_no, current_parent_title = number, right.strip().rstrip(".")
            main_clauses.append({
                "id": make_clause_id(number), "book": "FIDIC Red Book 2017", "edition": "2017",
                "source_form": "FIDIC 2017 Red", "parent_clause_no": number,
                "parent_clause_title": current_parent_title, "clause_no": number,
                "clause_title": current_parent_title, "level": "main-clause",
                "original_order": len(main_clauses), "source_text_origin": SOURCE_ORIGIN,
                "source_status": "source_text_loaded", "verification_status": "needs_pdf_verification",
                "parsing_status": "parsed", "qc_notes": [], "functional_mappings": [],
                "tags": [], "clause_elements": [],
            })
            continue
        if not current_parent_no or number.split(".", 1)[0] != current_parent_no:
            issues.append(f"Row {row_index}: {number} is not under the active parent clause")
        title = " ".join(part.strip() for part in label_tail.splitlines() if part.strip())
        if not right:
            omitted_empty_containers.append(number)
            continue
        qc_notes = []
        if not title: qc_notes.append("Sub-clause title is missing in corrected Word row.")
        if not right: qc_notes.append("Full text is empty in corrected Word row.")
        if "\ufffd" in title or "\ufffd" in right: qc_notes.append("Unicode replacement character detected.")
        clauses.append({
            "id": make_clause_id(number), "book": "FIDIC Red Book 2017", "edition": "2017",
            "source_form": "FIDIC 2017 Red", "parent_clause_no": current_parent_no,
            "parent_clause_title": current_parent_title, "clause_no": number,
            "clause_title": title, "level": "sub-clause", "original_order": len(clauses),
            "full_text": right, "paragraphs": paragraph_records(right),
            "literal_cross_references": literal_cross_references(right),
            "source_text_origin": SOURCE_ORIGIN, "source_status": "source_text_loaded",
            "verification_status": "needs_pdf_verification",
            "parsing_status": "needs_review" if qc_notes else "parsed", "qc_notes": qc_notes,
            "functional_mappings": [], "tags": [], "clause_elements": [],
        })
    for parent in main_clauses:
        parent["sub_clause_count"] = sum(x["parent_clause_no"] == parent["clause_no"] for x in clauses)
    numbers = [x["clause_no"] for x in clauses]
    qc = {
        "source_row_count": len(table.rows), "unparsed_row_issues": issues,
        "duplicate_sub_clause_numbers": sorted(n for n, count in Counter(numbers).items() if count > 1),
        "parent_mismatches": [x["clause_no"] for x in clauses if x["clause_no"].split(".", 1)[0] != x["parent_clause_no"]],
        "missing_number_records": [], "missing_title_records": [x["clause_no"] for x in clauses if not x["clause_title"]],
        "empty_full_text_records": [x["clause_no"] for x in clauses if not x["full_text"]],
        "needs_review_records": [x["clause_no"] for x in clauses if x["parsing_status"] == "needs_review"],
        "omitted_empty_container_rows": omitted_empty_containers,
        "corrected_docx_total_clause_rows": len(main_clauses) + len(clauses),
    }
    return main_clauses, clauses, qc


def build_source(path: Path) -> dict:
    main_clauses, clauses, qc = parse_docx(path)
    return {
        "schema_version": "1.1", "book": "FIDIC Red Book 2017", "edition": "2017",
        "source_file": path.as_posix(), "source_text_origin": SOURCE_ORIGIN,
        "source_status": "source_text_loaded", "verification_status": "needs_pdf_verification",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "source_sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "main_clause_count": len(main_clauses), "sub_clause_count": len(clauses),
        "main_clauses": main_clauses, "clauses": clauses, "quality_control": qc,
    }


def report(source: dict) -> str:
    qc = source["quality_control"]
    counts = "\n".join(f"| {x['clause_no']} | {x['sub_clause_count']} |" for x in source["main_clauses"])
    return f"""# FIDIC 2017 Red Corrected Word Import Report

## Source

- Corrected Word source: `{source['source_file']}`
- Source SHA-256: `{source['source_sha256']}`
- Extraction: python-docx, single two-column clause table
- Full wording source: corrected Word working copy

## Imported inventory

- Main clauses: **{source['main_clause_count']}**
- Sub-clause records: **{source['sub_clause_count']}**
- Source table rows including header: **{qc['source_row_count']}**

| Main clause | Sub-clause records |
| --- | ---: |
{counts}

## Controls

- Main sequence 1-21: **{[x['clause_no'] for x in source['main_clauses']] == [str(i) for i in range(1, 22)]}**
- Duplicate clause numbers: **{len(qc['duplicate_sub_clause_numbers'])}**
- Parent mismatches: **{len(qc['parent_mismatches'])}**
- Missing titles: **{len(qc['missing_title_records'])}**
- Empty full text: **{len(qc['empty_full_text_records'])}**
- Records needing parsing review: **{len(qc['needs_review_records'])}**

The importer regenerated the Clause Spine source layer from the corrected DOCX. It did not add mappings, tags, clause elements or legal analysis. PDF verification is performed separately.
"""


def main() -> int:
    parser = argparse.ArgumentParser(); parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path, required=True); parser.add_argument("--report", type=Path, required=True)
    parser.add_argument("--validate-only", action="store_true"); args = parser.parse_args()
    source = build_source(args.input); errors = validate_source_layer(source)
    print(json.dumps({"main_clause_count": source["main_clause_count"], "sub_clause_count": source["sub_clause_count"], "quality_control": source["quality_control"], "validation_errors": errors}, ensure_ascii=False, indent=2))
    if errors: return 1
    if not args.validate_only:
        args.output.parent.mkdir(parents=True, exist_ok=True); args.report.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(json.dumps(source, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        args.report.write_text(report(source), encoding="utf-8")
    return 0


if __name__ == "__main__": raise SystemExit(main())
